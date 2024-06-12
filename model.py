import pickle

import numpy as np
from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.neighbors import NearestCentroid
from sklearn.svm import SVC
import keras
from datetime import datetime
from functools import partial
from docx import Document
from docx.shared import Cm
from PIL import Image
import rasterio
import os
from multiprocessing import Pool
import random
from rasterio.transform import Affine
import xgboost as xgb
import matplotlib.pyplot as plt
from shapely.geometry import box

from MLClassifier import MLClassifier
from bbox import Bbox
from config import language, workspace
from dataset import load_dataset, RasterDataset
from report_tools import insert_table, insert_param_to_table, percent_format
from image import Sentinel2Image, load_image


def random_forest(
        train_features,
        train_labels,
        test_features,
        n_estimators,
        table
):
    insert_param_to_table(table, 'model_type', language['random_forest'])
    insert_param_to_table(table, 'n_estimators', n_estimators)
    rf = RandomForestClassifier(n_estimators=n_estimators)
    start = datetime.now()
    rf.fit(train_features, train_labels)
    end = datetime.now()
    insert_param_to_table(table, 'fitting_time', end - start)
    return rf, rf.predict(test_features)


def nearest_centroid(
        train_features,
        train_labels,
        test_features,
        table
):
    insert_param_to_table(table, 'model_type', language['nearest_centroid'])
    nc = NearestCentroid()
    start = datetime.now()
    nc.fit(train_features, train_labels)
    end = datetime.now()
    insert_param_to_table(table, 'fitting_time', end - start)
    return nc, nc.predict(test_features)


def svm(
        train_features,
        train_labels,
        test_features,
        table,
        gamma
):
    insert_param_to_table(table, 'model_type', language['svm'])
    m = SVC(gamma=gamma)
    start = datetime.now()
    m.fit(train_features, train_labels)
    end = datetime.now()
    insert_param_to_table(table, 'fitting_time', end - start)
    return m, m.predict(test_features)


def maximum_likelihood(
        train_features,
        train_labels,
        test_features,
        table
):
    insert_param_to_table(table, 'model_type', language['ml'])
    m = MLClassifier()
    start = datetime.now()
    m.fit(train_features, train_labels)
    end = datetime.now()
    insert_param_to_table(table, 'fitting_time', end - start)
    return m, m.predict(test_features)


def xgboost(
        train_features,
        train_labels,
        test_features,
        table,
        n_estimators=7,
        max_depth=4
):
    insert_param_to_table(table, 'model_type', language['xgboost'])
    insert_param_to_table(table, 'n_estimators', n_estimators)
    insert_param_to_table(table, 'max_depth', max_depth)
    m = xgb.XGBClassifier(n_estimators=n_estimators, max_depth=max_depth)
    start = datetime.now()
    m.fit(train_features, train_labels)
    end = datetime.now()
    insert_param_to_table(table, 'fitting_time', end - start)
    return m, m.predict(test_features)


def neural_network(
        train_features,
        train_labels,
        test_features,
        inputs,
        architecture,
        table,
        output_activation='softmax',
        optimizer=keras.optimizers.RMSprop(learning_rate=10**-9),
        loss='categorical_crossentropy',
        batch_size=32,
        epochs=10,
        loss_diff=1e-5
):
    insert_param_to_table(table, 'output_activation', output_activation)
    insert_param_to_table(table, 'optimizer', optimizer)
    insert_param_to_table(table, 'loss', loss),
    insert_param_to_table(table, 'batch_size', batch_size)
    insert_param_to_table(table, 'epochs', epochs)
    insert_param_to_table(table, 'loss_diff', loss_diff)
    num_classes = len(set(train_labels))
    train_labels = keras.utils.to_categorical(train_labels, num_classes).reshape(train_labels.shape[0], num_classes, 1)
    outputs = keras.layers.Dense(num_classes, activation=output_activation)(architecture)
    model = keras.Model(inputs=inputs, outputs=outputs)
    model.compile(optimizer=optimizer, loss=loss)
    loss = -1
    ploss = -2
    epochs_count = 0
    start = datetime.now()
    while abs(loss - ploss) > loss_diff:
        epochs_count += epochs
        ploss = loss
        model.fit(train_features, train_labels, batch_size=batch_size, epochs=epochs)
        loss = model.evaluate(train_features, train_labels)
    insert_param_to_table(table, 'fitting_time', datetime.now() - start)
    insert_param_to_table(table, 'epochs_count', epochs_count)
    predictions = model.predict(test_features)
    predictions = predictions.argmax(axis=1)
    return model, predictions


def perceptron(
        train_features,
        train_labels,
        test_features,
        hidden_layers,
        layer_size,
        table,
        activation='relu',
        output_activation='softmax',
        optimizer='rmsprop',
        loss='categorical_crossentropy',
        batch_size=32,
        epochs=10,
        loss_diff=1e-5
):
    inputs = keras.Input(shape=(train_features.shape[1]))
    insert_param_to_table(table, 'model_type', 'Perceptron')
    insert_param_to_table(table, 'hidden_layers', hidden_layers)
    insert_param_to_table(table, 'hidden_layer_size', layer_size)
    x = inputs
    for i in range(hidden_layers):
        x = keras.layers.Dense(layer_size, activation=activation)(x)
    return neural_network(
        train_features,
        train_labels,
        test_features,
        inputs,
        table=table,
        architecture=x,
        output_activation=output_activation,
        optimizer=optimizer,
        loss=loss,
        batch_size=batch_size,
        epochs=epochs,
        loss_diff=loss_diff
    )


def conv5(
        train_features,
        train_labels,
        test_features,
        first_conv_filters,
        second_conv_filters,
        layer_size,
        hidden_layers,
        table,
        activation='relu',
        output_activation='softmax',
        optimizer='rmsprop',
        loss='categorical_crossentropy',
        batch_size=32,
        epochs=10,
        loss_diff=1e-5
):
    inputs = keras.Input(shape=(train_features.shape[1:]))
    header_cells = table.rows[0].cells
    header_cells[0].text = language['parameter_name']
    header_cells[1].text = language['parameter_value']
    insert_param_to_table(table, 'model_type', 'Conv5')
    insert_param_to_table(table, 'first_conv_filters', first_conv_filters)
    insert_param_to_table(table, 'second_conv_filters', second_conv_filters)
    insert_param_to_table(table, 'hidden_layers', hidden_layers)
    insert_param_to_table(table, 'hidden_layer_size', layer_size)

    x = keras.layers.Conv2D(filters=first_conv_filters, kernel_size=(3, 3), activation=activation)(inputs)
    x = keras.layers.Conv2D(filters=second_conv_filters, kernel_size=(3, 3), activation=activation)(x)
    x = keras.layers.GlobalAveragePooling2D()(x)
    for i in range(hidden_layers):
        x = keras.layers.Dense(layer_size, activation=activation)(x)
    return neural_network(
        train_features,
        train_labels,
        test_features,
        inputs,
        table=table,
        architecture=x,
        output_activation=output_activation,
        optimizer=optimizer,
        loss=loss,
        batch_size=batch_size,
        epochs=epochs,
        loss_diff=loss_diff
    )


def conv7(
        train_features,
        train_labels,
        test_features,
        first_conv_filters,
        second_conv_filters,
        third_conv_filters,
        layer_size,
        hidden_layers,
        table,
        activation='relu',
        output_activation='softmax',
        optimizer='rmsprop',
        loss='categorical_crossentropy',
        batch_size=32,
        epochs=10,
        loss_diff=1e-5
):
    inputs = keras.Input(shape=(train_features.shape[1:]))
    header_cells = table.rows[0].cells
    header_cells[0].text = language['parameter_name']
    header_cells[1].text = language['parameter_value']
    insert_param_to_table(table, 'model_type', 'Conv7')
    insert_param_to_table(table, 'first_conv_filters', first_conv_filters)
    insert_param_to_table(table, 'second_conv_filters', second_conv_filters)
    insert_param_to_table(table, 'third_conv_filters', third_conv_filters)
    insert_param_to_table(table, 'hidden_layers', hidden_layers)
    insert_param_to_table(table, 'hidden_layer_size', layer_size)

    x = keras.layers.Conv2D(filters=first_conv_filters, kernel_size=(3, 3), activation=activation)(inputs)
    x = keras.layers.Conv2D(filters=second_conv_filters, kernel_size=(3, 3), activation=activation)(x)
    x = keras.layers.Conv2D(filters=third_conv_filters, kernel_size=(3, 3), activation=activation)(x)
    x = keras.layers.GlobalAveragePooling2D()(x)
    for i in range(hidden_layers):
        x = keras.layers.Dense(layer_size, activation=activation)(x)
    return neural_network(
        train_features,
        train_labels,
        test_features,
        inputs,
        table=table,
        architecture=x,
        output_activation=output_activation,
        optimizer=optimizer,
        loss=loss,
        batch_size=batch_size,
        epochs=epochs,
        loss_diff=loss_diff
    )


def spp_conv(
    train_features,
    train_labels,
    test_features,
    table,
    output_activation='softmax',
    optimizer='rmsprop',
    loss='categorical_crossentropy',
    batch_size=32,
    epochs=10,
    loss_diff=1e-5
):
    num_classes = len(set(train_labels))
    train_features_by_size = {}
    train_labels_by_size = {}
    for i in range(len(train_features)):
        if train_features[i][0] in train_features_by_size:
            train_features_by_size[train_features[i][0]].append(train_features[i][1])
            train_labels_by_size[train_features[i][0]].append(train_labels[i])
        else:
            train_features_by_size[train_features[i][0]] = [train_features[i][1]]
            train_labels_by_size[train_features[i][0]] = [train_labels[i]]
    for key in train_labels_by_size:
        train_labels_by_size[key] = keras.utils.to_categorical(
            train_labels_by_size[key],
            num_classes
        ).reshape(
            len(train_labels_by_size[key]), 2, 1
        )

    model = keras.Sequential()

    # uses theano ordering. Note that we leave the image size as None to allow multiple image sizes
    model.add(keras.layers.Convolution2D(32, 3, 3, padding='same', input_shape=(None, None, 3)))
    model.add(keras.layers.Activation('relu'))
    model.add(keras.layers.Convolution2D(32, 3, 3))
    model.add(keras.layers.Activation('relu'))
    model.add(keras.layers.MaxPooling2D(pool_size=(2, 2)))
    model.add(keras.layers.Convolution2D(64, 3, 3, padding='same'))
    model.add(keras.layers.Activation('relu'))
    model.add(keras.layers.Convolution2D(64, 3, 3))
    model.add(keras.layers.Activation('relu'))
    # model.add(SpatialPyramidPooling([1, 2, 4]))
    model.add(keras.layers.GlobalAveragePooling2D())
    model.add(keras.layers.Dense(num_classes))
    model.add(keras.layers.Activation(output_activation))

    model.compile(optimizer=optimizer, loss=loss)
    loss = -1
    ploss = -2
    epochs_count = 0
    start = datetime.now()
    cur_idx = 0
    while abs(loss - ploss) > loss_diff:
        ploss = loss
        for key in train_features_by_size:
            for i in range(0, len(train_features_by_size[key]), 16):
                with Pool() as pool:
                    loaded_features = pool.starmap(load_image, [(el,) for el in train_features_by_size[key][i:i + 16]])
                key = train_features[cur_idx][0]
                model.fit(np.array(loaded_features), np.array(train_labels_by_size[key][i:i + 16]))
            epochs_count += 1
        prediction = []
        answers = []
        for key in train_features_by_size:
            for i in range(0, len(train_features_by_size[key]), 16):
                with Pool() as pool:
                    loaded_features = pool.starmap(load_image, [(el,) for el in train_features_by_size[key][i:i + 16]])
                prediction += model.predict(np.array(loaded_features)).tolist()
            answers += train_labels_by_size[key].tolist()
        prediction = np.array(prediction)
        prediction = prediction[:, 0] > prediction[:, 1]
        answers = np.array(answers)[:, 0]
        loss = np.sum(prediction == answers)/prediction.shape[0]
    predictions = []
    for path in test_features:
        with rasterio.open(path) as image:
            predictions += model.predict([np.block([[[image.read(index)]] for index in image.indexes])])
    predictions = np.array(predictions).argmax(axis=1)
    return model, predictions


def confusion_matrix(labels, predictions, doc):
    labels_set = list(set(labels))
    conf_matrix = np.zeros((len(labels_set), len(labels_set)), dtype=np.int32)
    for i in range(len(labels_set)):
        for j in range(len(labels_set)):
            conf_matrix[i, j] = np.sum((labels == labels_set[j]) & (predictions == labels_set[i]))
    accuracy = np.array([conf_matrix[i, i] for i in range(conf_matrix.shape[1])]).sum() / labels.shape[0]
    precision = np.array([conf_matrix[i, i] / conf_matrix[i, :].sum() for i in range(conf_matrix.shape[0])])
    recall = np.array([conf_matrix[i, i] / conf_matrix[:, i].sum() for i in range(conf_matrix.shape[1])])
    f1score = 2 / (1 / precision + 1 / recall)

    doc.add_heading(language['confusion_matrix'], level=1)
    table = doc.add_table(rows=conf_matrix.shape[0] + 4, cols=conf_matrix.shape[0] + 3)
    header_cells = table.rows[0].cells
    header_cells[0].text = language['test\\training']
    for i in range(len(labels_set)):
        header_cells[i + 1].text = str(labels_set[i])
    header_cells[len(labels_set) + 1].text = language['sum']
    header_cells[len(labels_set) + 2].text = language['precision']

    for i in range(len(labels_set)):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = str(labels_set[i])
        for j in range(len(labels_set)):
            row_cells[j + 1].text = str(conf_matrix[i, j])
        row_cells[len(labels_set) + 1].text = str(conf_matrix[i, :].sum())
        row_cells[len(labels_set) + 2].text = percent_format(precision[i])

    row_cells = table.rows[len(labels_set) + 1].cells
    row_cells[0].text = language['sum']
    for i in range(len(labels_set)):
        row_cells[i + 1].text = str(conf_matrix[:, i].sum())
    row_cells[len(labels_set) + 1].text = str(conf_matrix.sum())

    row_cells = table.rows[len(labels_set) + 2].cells
    row_cells[0].text = language['recall']
    for i in range(len(labels_set)):
        row_cells[i + 1].text = percent_format(recall[i])
    row_cells[-1].text = f"{language['accuracy']}: {percent_format(accuracy)}"

    row_cells = table.rows[len(labels_set) + 3].cells
    row_cells[0].text = language['f1']
    for i in range(len(labels_set)):
        row_cells[i + 1].text = percent_format(f1score[i])
    p_0 = accuracy * labels.shape[0]
    p_e = np.array([conf_matrix[i, :].sum() * conf_matrix[:, i].sum() for i in range(conf_matrix.shape[0])]).sum()
    kappa = (labels.shape[0] * p_0 - p_e) / (labels.shape[0] ** 2 - p_e)
    row_cells[-1].text = f"{language['kappa']}: {percent_format(kappa)}"

    return accuracy, kappa


class Model:
    def __init__(self, name, classifier, normalize=False, pca_decomposition=False, pca_explained=0.99):
        self.name = name
        self.classifier = classifier
        self.fitted_model = None
        self.normalize = normalize
        self.normalizer = StandardScaler()
        self.pca_decomposition = pca_decomposition
        self.pca = PCA()
        self.pca_components = -1
        self.pca_explained = pca_explained
        self.is_conv = None
        self.window_size = None

    def learn(
            self,
            dataset_name,
            report_name,
            test_size=0.2
    ):
        doc = Document()
        doc.add_heading(f"{language['model_heading']}: {self.name}", level=0)
        table = insert_table(doc, 'learning_parameters')

        dataset = load_dataset(dataset_name)
        indices = np.array(range(len(dataset.features)))
        np.random.shuffle(indices)
        features = dataset.features[indices]
        labels = dataset.labels[indices]
        insert_param_to_table(table, 'dataset_name', dataset_name)
        insert_param_to_table(table, 'features_count', features.shape[0])
        insert_param_to_table(table, 'features_params', features.shape[-1])
        insert_param_to_table(table, 'normalization', self.normalize)
        self.is_conv = dataset.is_conv
        self.window_size = dataset.window_size
        if self.normalize:
            if dataset.is_conv:
                self.normalizer.fit(features[:, dataset.window_size, dataset.window_size])
                features = self.normalizer.transform(
                    features.reshape(
                        features.shape[0] * features.shape[1] * features.shape[2],
                        features.shape[-1]
                    )
                ).reshape(
                    features.shape[0],
                    features.shape[1],
                    features.shape[2],
                    features.shape[-1]
                )
            else:
                features = self.normalizer.fit_transform(features)
            insert_param_to_table(table, 'pca', self.pca_decomposition)
            if self.pca_decomposition:
                insert_param_to_table(table, 'pca_explained_threshold', percent_format(self.pca_explained))
                if dataset.is_conv:
                    self.pca.fit(features[:, dataset.window_size, dataset.window_size])
                    features = self.pca.transform(
                        features.reshape(
                            features.shape[0] * features.shape[1] * features.shape[2],
                            features.shape[-1]
                        )
                    ).reshape(
                        features.shape[0],
                        features.shape[1],
                        features.shape[2],
                        features.shape[-1]
                    )
                else:
                    features = self.pca.fit_transform(features)
                explained = 0.0
                for i in range(len(self.pca.explained_variance_ratio_)):
                    explained += self.pca.explained_variance_ratio_[i]
                    if explained >= self.pca_explained:
                        self.pca_components = i + 1
                        insert_param_to_table(table, 'pca_arguments', self.pca_components)
                        insert_param_to_table(table, 'pca_explained_final', percent_format(explained))
                        if dataset.is_conv:
                            features = features[:, :, :, :self.pca_components]
                        else:
                            features = features[:, :self.pca_components]
                        break

        insert_param_to_table(table, 'test_ratio', percent_format(test_size))
        if test_size > 0:
            train_features, test_features, train_labels, test_labels = train_test_split(
                features,
                labels,
                test_size=test_size
            )
        else:
            train_features = features
            test_features = features
            train_labels = labels
            test_labels = labels
        insert_param_to_table(table, 'training_count', train_features.shape[0])
        insert_param_to_table(table, 'test_count', test_features.shape[0])

        table = insert_table(doc, 'model_parameters')

        self.fitted_model, predictions = self.classifier(
            train_features=train_features,
            train_labels=train_labels,
            test_features=test_features,
            table=table
        )
        accuracy, kappa = confusion_matrix(test_labels, predictions, doc)

        doc.save(os.path.join(workspace, "reports", "learning", f"{report_name}.docx"))
        return accuracy, kappa

    def classify_image(
            self,
            image_name,
            output_file,
            bbox,
            apply_cloud_mask=None,
            invalid_value=255
    ):
        image = load_image(image_name)
        image.unzip()
        temp_path = os.path.join(workspace, "temp", image.identifier)
        layers = [
            el for el in os.listdir(temp_path)
            if '60m' not in el and 'B01' not in el and el[-4:] == '.jp2'
        ]
        features = []
        first = bbox is not None
        for file_name in layers:
            with (rasterio.open(os.path.join(temp_path, file_name)) as channel):
                if first:
                    first = False
                    bbox = bbox.to_crs(channel.crs)
                    bbox.adjust_to_grid(60)
                if bbox is None:
                    data, transform = rasterio.mask.mask(channel, [box(*channel.bounds)], crop=True)
                else:
                    data, transform = rasterio.mask.mask(channel, [bbox.as_shapely()], crop=True)
                if '10m' in file_name:
                    profile = channel.profile
                    profile.update(transform=transform)
                elif '20m' in file_name:
                    if 'B01' in file_name:
                        continue
                    data2 = np.zeros((1, data.shape[1] * 2, data.shape[2] * 2), dtype=data.dtype)
                    for i in range(2):
                        for j in range(2):
                            data2[:, i::2, j::2] = data
                    data = data2
                elif '60m' in file_name:
                    continue
                features.append(data[0])
        features = np.array(features)
        if bbox is not None:
            profile.update(
                width=features.shape[2],
                height=features.shape[1]
            )
        original_shape = (features.shape[1], features.shape[2])
        if self.window_size == 0:
            features = features.reshape(features.shape[0], features.shape[1]*features.shape[2]).transpose()
        else:
            features2 = np.zeros(
                (
                    features.shape[0],
                    features.shape[1] + 2 * self.window_size,
                    features.shape[2] + 2 * self.window_size
                ),
                dtype=features.dtype
            )
            features2[
                :,
                self.window_size:features.shape[1] + self.window_size,
                self.window_size:features.shape[2] + self.window_size
            ] = features
            features3 = np.zeros(
                (
                    features.shape[0],
                    self.window_size * 2 + 1,
                    self.window_size * 2 + 1,
                    features.shape[1],
                    features.shape[2],
                ), dtype=features.dtype)
            for i in range(2 * self.window_size + 1):
                for j in range(2 * self.window_size + 1):
                    features3[:, i, j] = features2[:, i:features.shape[1] + i, j:features.shape[2] + j]
            if self.is_conv:
                features = features3.reshape((
                    features.shape[0],
                    self.window_size * 2 + 1,
                    self.window_size * 2 + 1,
                    features.shape[1]*features.shape[2]
                )).transpose(3, 1, 2, 0)
            else:
                features = features3.reshape((
                    features.shape[0] * (self.window_size * 2 + 1)**2,
                    features.shape[1] * features.shape[2]
                )).transpose()
        if self.normalize:
            if self.is_conv:
                features = self.normalizer.transform(
                    features.reshape(
                        features.shape[0] * features.shape[1] * features.shape[2],
                        features.shape[-1]
                    )
                ).reshape(
                    features.shape[0],
                    features.shape[1],
                    features.shape[2],
                    features.shape[-1]
                )
            else:
                features = self.normalizer.transform(features)
            if self.pca_decomposition:
                features = self.pca.transform(features)[:, :self.pca_components]
        predictions = self.fitted_model.predict(features)
        if len(predictions.shape) > 1:
            predictions = predictions.argmax(axis=1)
        predictions = predictions.reshape(original_shape[0], original_shape[1])
        if isinstance(apply_cloud_mask, list):
            with rasterio.open(
                os.path.join(
                    workspace,
                    "cloud_masks",
                    f"{image.name.rsplit('_', 1)[0].replace('MSIL2A', 'MSIL1C')}.tif"
                )
            ) as mask:
                if bbox is None:
                    masked, transform = rasterio.mask.mask(mask, [box(*channel.bounds)], crop=True)
                else:
                    masked, transform = rasterio.mask.mask(mask, [bbox.as_shapely()], crop=True)
                masked2 = np.zeros((1, masked.shape[1] * 6, masked.shape[2] * 6), dtype=masked.dtype)
                for k in range(6):
                    for j in range(6):
                        masked2[:, k::6, j::6] = masked
                masked = masked2
                for el in apply_cloud_mask:
                    masked &= (masked[0] & el) == 0
            predictions[masked] = invalid_value
        profile.update(
            dtype=rasterio.uint8
        )
        with rasterio.open(os.path.join(workspace, "classImages", f"{output_file}.tif"), "w", **profile) as dst:
            dst.write(predictions.astype(rasterio.uint8), 1)

    def classify_dataset(
            self,
            dataset_name,
            report_name
    ):
        doc = Document()
        doc.add_heading(f"{language['model_heading']}: {self.name}", level=0)
        table = insert_table(doc, 'learning_parameters')

        dataset = load_dataset(dataset_name)
        features = dataset.features
        labels = dataset.labels
        insert_param_to_table(table, 'dataset_name', dataset_name)
        insert_param_to_table(table, 'features_count', features.shape[0])
        insert_param_to_table(table, 'features_params', features.shape[-1])
        insert_param_to_table(table, 'normalization', self.normalize)
        if self.normalize:
            if dataset.is_conv:
                features = self.normalizer.transform(
                    features.reshape(
                        features.shape[0] * features.shape[1] * features.shape[2],
                        features.shape[-1]
                    )
                ).reshape(
                    features.shape[0],
                    features.shape[1],
                    features.shape[2],
                    features.shape[-1]
                )
            else:
                features = self.normalizer.transform(features)
            insert_param_to_table(table, 'pca', self.pca_decomposition)
            if self.pca_decomposition:
                insert_param_to_table(table, 'pca_explained_threshold', percent_format(self.pca_explained))
                if dataset.is_conv:
                    features = self.pca.transform(
                        features.reshape(
                            features.shape[0] * features.shape[1] * features.shape[2],
                            features.shape[-1]
                        )
                    )[:, :self.pca_components].reshape(
                        features.shape[0],
                        features.shape[1],
                        features.shape[2],
                        self.pca_components
                    )
                else:
                    features = self.pca.transform(features)[:, :self.pca_components]
                insert_param_to_table(table, 'pca_arguments', self.pca_components)

        predictions = self.fitted_model.predict(features)
        if len(predictions.shape) > 1:
            predictions = predictions.argmax(axis=1)
        accuracy, kappa = confusion_matrix(labels, predictions, doc)

        doc.save(os.path.join(workspace, "reports", "classifying", f"{report_name}.docx"))
        return accuracy, kappa

    def save(self):
        pickle.dump(self, open(os.path.join(workspace, "models", self.name), "wb"))


def transform_image(transformer, dataset_name, path):
    final_path = os.path.join(workspace, "parcels", dataset_name, "unnormalized", path[0], path[1])
    with rasterio.open(final_path) as image:
        data = np.block([[[image.read(index)]] for index in image.indexes]
                        ).reshape((len(image.indexes), image.width * image.height)).transpose()
        data = transformer.transform(data).transpose(
        ).reshape((len(image.indexes), image.height, image.width))
        dest_profile = image.profile
        dest_profile.update(dtype=rasterio.float32)
        level_path = os.path.join(workspace, "parcels", dataset_name, "normalized", path[0])
        if not os.path.exists(level_path):
            os.mkdir(level_path)
        image_path = os.path.join(level_path, path[1])
        with rasterio.open(
                image_path,
                mode="w",
                **dest_profile
        ) as dest:
            dest.write(data)


class CadastralModel:
    def __init__(self, name, classifier, normalize=True):
        self.name = name
        self.classifier = classifier
        self.fitted_model = None
        self.normalize = normalize
        self.normalizer = StandardScaler()

    def learn(
            self,
            dataset_name,
            report_name,
            test_size=0.2
    ):
        doc = Document()
        doc.add_heading(f"{language['model_heading']}: {self.name}", level=0)
        table = insert_table(doc, 'learning_parameters')

        cadastral_dataset_path = os.path.join(workspace, "parcels", dataset_name, "unnormalized")
        feature_list = []
        for el in os.listdir(cadastral_dataset_path):
            for el2 in os.listdir(os.path.join(cadastral_dataset_path, el)):
                feature_list.append([el, el2])
        insert_param_to_table(table, 'dataset_name', dataset_name)
        insert_param_to_table(table, 'features_count', len(feature_list))
        with rasterio.open(os.path.join(cadastral_dataset_path, feature_list[0][0], feature_list[0][1])) as image:
            insert_param_to_table(table, 'features_params', len(image.indexes))
        insert_param_to_table(table, 'normalization', self.normalize)
        if self.normalize:
            i = 0
            for path in feature_list:
                i += 1
                if i % 100 == 0:
                    print(i/len(feature_list))
                    break
                with rasterio.open(os.path.join(cadastral_dataset_path, path[0], path[1])) as image:
                    data = np.block([[[image.read(index)]] for index in image.indexes]
                                    ).reshape((len(image.indexes), image.width * image.height)).transpose()
                    self.normalizer.partial_fit(data[data[:, 0] > 0, :])
            normalized_path = os.path.join(workspace, "temp", "parcels", dataset_name, "normalized")
            if not os.path.exists(normalized_path):
                os.mkdir(normalized_path)
            with Pool() as pool:
                pool.starmap(
                    transform_image,
                    [
                        (
                            self.normalizer,
                            dataset_name,
                            path
                        ) for path in feature_list
                    ]
                )
            feature_list = [[el[0], os.path.join(normalized_path, el[0], el[1])] for el in feature_list]
        else:
            feature_list = [[el[0], os.path.join(cadastral_dataset_path, el[0], el[1])] for el in feature_list]
        insert_param_to_table(table, 'test_ratio', percent_format(test_size))
        random.shuffle(feature_list)
        labels = [int(el[1].rsplit("_", 1)[1].split(".", 1)[0]) for el in feature_list]
        border_index = int(len(feature_list) * test_size)
        train_features = feature_list[border_index:]
        train_labels = np.array(labels[border_index:])
        test_features = feature_list[:border_index]
        test_labels = np.array(labels[:border_index])
        insert_param_to_table(table, 'training_count', len(train_features))
        insert_param_to_table(table, 'test_count', len(test_features))

        table = insert_table(doc, 'model_parameters')

        self.fitted_model, predictions = self.classifier(
            train_features=train_features,
            train_labels=train_labels,
            test_features=test_features,
            table=table
        )
        true_negatives = np.sum((predictions < 0.5) & (test_labels == 0))
        false_negatives = np.sum((predictions < 0.5) & (test_labels == 1))
        false_positives = np.sum((predictions >= 0.5) & (test_labels == 0))
        true_positives = np.sum((predictions >= 0.5) & (test_labels == 1))

        accuracy = (true_positives + true_negatives) / test_labels.shape[0]

        try:
            precision0 = true_negatives / (true_negatives + false_negatives)
        except:
            precision0 = "NaN"
        try:
            recall0 = true_negatives / (true_negatives + false_positives)
        except:
            recall0 = "NaN"
        try:
            precision1 = true_positives / (false_positives + true_positives)
        except:
            precision1 = "NaN"
        try:
            recall1 = true_positives / (false_negatives + true_positives)
        except:
            recall1 = "NaN"

        try:
            f10 = 2 / (1 / precision0 + 1 / recall0)
        except:
            f10 = "NaN"
        try:
            f11 = 2 / (1 / precision1 + 1 / recall1)
        except:
            f11 = "NaN"

        try:
            mean_f1 = (f10 + f11) / 2
        except:
            mean_f1 = "NaN"

        doc.add_heading(language['confusion_matrix'], level=1)
        table = doc.add_table(rows=5, cols=4)
        header_cells = table.rows[0].cells
        header_cells[0].text = language['test\\training']
        header_cells[1].text = "0"
        header_cells[2].text = "1"
        header_cells[3].text = language['precision']

        row_cells = table.rows[1].cells
        row_cells[0].text = "0"
        row_cells[1].text = str(true_negatives)
        row_cells[2].text = str(false_negatives)
        row_cells[3].text = percent_format(precision0)

        row_cells = table.rows[2].cells
        row_cells[0].text = "1"
        row_cells[1].text = str(false_positives)
        row_cells[2].text = str(true_positives)
        row_cells[3].text = percent_format(precision1)

        row_cells = table.rows[3].cells
        row_cells[0].text = language['recall']
        row_cells[1].text = percent_format(recall0)
        row_cells[2].text = percent_format(recall1)
        row_cells[3].text = f"{language['accuracy']}: {percent_format(accuracy)}"

        row_cells = table.rows[4].cells
        row_cells[0].text = language['f1']
        row_cells[1].text = percent_format(f10)
        row_cells[2].text = percent_format(f11)
        row_cells[3].text = f"{language['f1_mean']}: {percent_format(mean_f1)}"
        doc.save(os.path.join(workspace, "reports", "learning", f"{report_name}.docx"))
        return accuracy, mean_f1


def load_model(name):
    return pickle.load(open(os.path.join(workspace, "models", name), "rb"))
