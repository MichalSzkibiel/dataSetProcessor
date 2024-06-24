import shutil

import pandas as pd
import rasterio
import rasterio.mask
import rasterio.merge
import os
import numpy as np
import pickle
import geopandas as gpd
from docx import Document
from multiprocessing import Process as Thread, Pool
import math
import matplotlib.pyplot as plt
import random

from shapely.geometry import box

from feature_downloader import osm_data
from bbox import Bbox, from_shapely_bounds
from image import Sentinel2Image, load_image, download_image
from rasterizer import process_polygon
from config import language, workspace
from report_tools import insert_table, insert_param_to_table


class MergedDataset:
    def __init__(
            self,
            dataset_name,
            datasets
    ):
        if isinstance(dataset_name, str):
            self.dataset_name = dataset_name
        else:
            Exception("dataset_name must be str")
        correct_type = True
        if isinstance(datasets, list):
            for el in datasets:
                if not isinstance(el, PartialDataset):
                    correct_type = False
        if not isinstance(datasets, list) and not correct_type:
            Exception("datasets must be list of PartialDataset")
        self.sentinel2image = datasets[0].sentinel2image
        self.window_size = datasets[0].window_size
        self.is_conv = datasets[0].is_conv
        self.ignore_off_scope = datasets[0].ignore_off_scope
        self.apply_rotations_and_reflections = datasets[0].apply_rotations_and_reflections

        self.features = []
        self.labels = []
        for el in datasets:
            self.features += el.features.tolist()
            self.labels += el.labels.tolist()
        self.features = np.array(self.features)
        self.labels = np.array(self.labels)

        self.write_report()

    def save_dataset(self):
        pickle.dump(
            self,
            open(os.path.join(workspace, "datasets", self.dataset_name), "wb"),
            protocol=pickle.HIGHEST_PROTOCOL
        )

    def write_report(self):
        doc = Document()
        doc.add_heading(f"{language['dataset_heading']}: {self.dataset_name}", level=0)
        table = insert_table(doc, 'dataset_parameters')

        insert_param_to_table(table, 'sentinel2_name', self.sentinel2image.name)
        insert_param_to_table(table, 'sentinel2_date', self.sentinel2image.date_time)
        insert_param_to_table(table, 'window_size', self.window_size * 2 + 1)
        if self.window_size > 0:
            insert_param_to_table(table, 'ignore_off_scope', self.ignore_off_scope)
            insert_param_to_table(table, 'is_conv', self.is_conv)
            insert_param_to_table(table, 'rotations_and_reflections', self.apply_rotations_and_reflections)

        insert_param_to_table(table, 'features_count', self.features.shape[0])
        insert_param_to_table(table, 'features_params', self.features.shape[-1])

        doc.add_heading(language['labels_count'], level=1)
        table = doc.add_table(rows=1, cols=2)
        header_cells = table.rows[0].cells
        header_cells[0].text = language['class']
        header_cells[1].text = language['count']

        for i in range(self.labels.max() + 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(np.sum(self.labels == i))

        doc.save(os.path.join(workspace, "reports", "datasets", f"{self.dataset_name}.docx"))


class PartialDataset:
    def __init__(
            self,
            dataset,
            subset_id,
            features,
            labels
    ):
        self.dataset_name = dataset.dataset_name
        self.sentinel2image = dataset.sentinel2image
        self.window_size = dataset.window_size
        self.is_conv = dataset.is_conv
        self.apply_rotations_and_reflections = dataset.apply_rotations_and_reflections
        if isinstance(subset_id, str):
            self.subset_id = subset_id
        else:
            Exception("dataset_name must be str")
        self.features = features
        print(self.features.shape)
        self.labels = labels

        self.write_report()

    def save_dataset(self):
        parts_dir = os.path.join(workspace, "datasets", f"{self.dataset_name}_parts")
        if not os.path.exists(parts_dir):
            os.mkdir(parts_dir)
        if os.path.exists(os.path.join(parts_dir, self.subset_id)):
            os.remove(os.path.join(parts_dir, self.subset_id))
        pickle.dump(self, open(os.path.join(parts_dir, self.subset_id), "wb"),
                    protocol=pickle.HIGHEST_PROTOCOL)

    def write_report(self):
        doc = Document()
        doc.add_heading(f"{language['dataset_heading']}: {self.dataset_name} {self.subset_id}", level=0)
        table = insert_table(doc, 'dataset_parameters')

        insert_param_to_table(table, 'sentinel2_name', [image.name for image in self.sentinel2image])
        insert_param_to_table(table, 'sentinel2_date', self.sentinel2image[0].date_time)
        insert_param_to_table(table, 'window_size', self.window_size * 2 + 1)
        if self.window_size > 0:
            insert_param_to_table(table, 'is_conv', self.is_conv)
            insert_param_to_table(table, 'rotations_and_reflections', self.apply_rotations_and_reflections)

        insert_param_to_table(table, 'features_count', self.features.shape[0])
        insert_param_to_table(table, 'features_params', self.features.shape[-1])

        doc.add_heading(language['labels_count'], level=1)
        table = doc.add_table(rows=1, cols=2)
        header_cells = table.rows[0].cells
        header_cells[0].text = language['class']
        header_cells[1].text = language['count']

        for i in range(self.labels.max() + 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(np.sum(self.labels == i))

        parts_dir = os.path.join(workspace, "reports", "datasets", f"{self.dataset_name}_parts")
        if not os.path.exists(parts_dir):
            os.mkdir(parts_dir)
        doc.save(os.path.join(parts_dir, f"{self.subset_id}.docx"))


class VectorDataset:
    def __init__(
            self,
            dataset_name,
            label_source,
            class_column,
            sentinel2image,
            apply_cloud_mask=False,
            window_size=1,
            is_conv=False,
            apply_rotations_and_reflections=False
    ):
        if isinstance(dataset_name, str):
            self.dataset_name = dataset_name
        else:
            Exception("dataset_name must be str")
        self.label_source = label_source
        self.bbox = from_shapely_bounds(self.label_source.total_bounds, label_source.crs)
        self.class_column = class_column
        if isinstance(sentinel2image, Sentinel2Image):
            self.sentinel2image = [sentinel2image]
        elif isinstance(sentinel2image, list):
            self.sentinel2image = sentinel2image
        else:
            Exception("identifier must be str")
        if isinstance(apply_cloud_mask, list):
            self.apply_cloud_mask = 0
            for el in apply_cloud_mask:
                self.apply_cloud_mask += el
        elif not apply_cloud_mask:
            self.apply_cloud_mask = apply_cloud_mask
        else:
            Exception("apply_cloud_mask must be False or list")
        if isinstance(window_size, int):
            self.window_size = window_size
        else:
            Exception("windows_size must be int")
        if isinstance(is_conv, bool):
            self.is_conv = is_conv
        else:
            Exception("ignore_off_scope must be bool")
        if self.window_size == 0:
            self.apply_rotations_and_reflections = False
        elif isinstance(apply_rotations_and_reflections, bool):
            self.apply_rotations_and_reflections = apply_rotations_and_reflections
        else:
            Exception("ignore_off_scope must be bool")

        snippets = []
        for i, row in label_source.iterrows():
            geom_bbox = Bbox(
                minx=(row.geometry.bounds[0] - self.window_size * 10) // 60 * 60,
                miny=(row.geometry.bounds[1] - self.window_size * 10) // 60 * 60,
                maxx=math.ceil((row.geometry.bounds[2] + self.window_size * 10) / 60) * 60,
                maxy=math.ceil((row.geometry.bounds[3] + self.window_size * 10) / 60) * 60,
                crs=label_source.crs
            )
            polygon_mask = process_polygon(row.geometry, geom_bbox, (10, 10))
            snippets.append([
                geom_bbox,
                polygon_mask,
                row[class_column],
                np.zeros((10, polygon_mask.shape[0], polygon_mask.shape[1]), dtype=np.uint16)
            ])
        self.features = []
        self.labels = []
        if isinstance(self.apply_cloud_mask, int):
            for image in self.sentinel2image:
                with rasterio.open(
                        os.path.join(
                            workspace,
                            "cloud_masks",
                            f"{image.name.rsplit('_', 1)[0].replace('MSIL2A', 'MSIL1C')}.tif"
                        )) as mask:
                    for i in range(len(snippets)):
                        try:
                            masked, transform = rasterio.mask.mask(mask, [snippets[i][0].as_shapely()], crop=True)
                        except:
                            continue
                        xshift = int((transform[2] - snippets[i][0].minx) // 10)
                        yshift = int((snippets[i][0].maxy - transform[5]) // 10)
                        masked2 = np.zeros((1, masked.shape[1] * 6, masked.shape[2] * 6), dtype=masked.dtype)
                        for k in range(6):
                            for j in range(6):
                                masked2[:, k::6, j::6] = masked
                        masked = masked2
                        snippets[i][1][
                            yshift:masked.shape[1] + yshift,
                            xshift:masked.shape[2] + xshift
                        ] = (masked[0] & self.apply_cloud_mask) == 0
        for i in range(len(snippets) - 1, -1, -1):
            if np.max(snippets[i][1]) == 0:
                snippets = snippets[:i] + snippets[i + 1:]
        for image in self.sentinel2image:
            image.unzip()
            idx = 0
            layers = [
                el for el in os.listdir(os.path.join(workspace, "temp", image.identifier))
                if '60m' not in el and 'B01' not in el and el[-4:] == '.jp2'
            ]
            print(layers)
            for file_name in layers:
                with rasterio.open(os.path.join(workspace, "temp", image.identifier, file_name)) as channel:
                    for i in range(len(snippets)):
                        try:
                            masked, transform = rasterio.mask.mask(channel, [snippets[i][0].as_shapely()], crop=True)
                        except:
                            continue
                        xshift = int((transform[2] - snippets[i][0].minx) // 10)
                        yshift = int((snippets[i][0].maxy - transform[5]) // 10)
                        if '20m' in file_name:
                            if 'B01' in file_name:
                                continue
                            masked2 = np.zeros((1, masked.shape[1] * 2, masked.shape[2] * 2), dtype=masked.dtype)
                            for k in range(2):
                                for j in range(2):
                                    masked2[:, k::2, j::2] = masked
                            masked = masked2
                        elif '60m' in file_name:
                            continue

                        snippets[i][3][idx, yshift:masked.shape[1] + yshift, xshift:masked.shape[2] + xshift] = masked[
                            0]
                    idx += 1
        for el in snippets:
            for i in range(self.window_size, el[3].shape[1] - self.window_size):
                for j in range(self.window_size, el[3].shape[2] - self.window_size):
                    if el[1][i, j]:
                        self.features.append(
                            el[3][
                            :,
                            i - self.window_size:i + self.window_size + 1,
                            j - self.window_size:j + self.window_size + 1
                            ])
                        self.labels.append(el[2])
        self.features = np.array(self.features)
        self.labels = np.array(self.labels)
        del label_source
        del snippets
        print(set(self.labels))
        print(self.labels.shape)
        print(self.features.shape)
        if self.apply_rotations_and_reflections:
            self.features = rotations_and_reflections(self.features)
            self.labels = np.block([
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels
            ])
        if not self.is_conv:
            self.features = self.features.reshape(
                self.features.shape[0],
                (self.window_size * 2 + 1) ** 2 * 10
            )
        if self.is_conv:
            self.features = self.features.transpose(0, 2, 3, 1)
        print(self.features.shape)
        self.write_report()

    def save_dataset(self):
        pickle.dump(
            self,
            open(os.path.join(workspace, "datasets", self.dataset_name), "wb"),
            protocol=pickle.HIGHEST_PROTOCOL
        )

    def split_dataset(self, parts, elements):
        indices = np.arange(len(self.labels))
        labels_indices = {el: indices[self.labels == el] for el in set(self.labels)}
        for label in labels_indices:
            np.random.shuffle(labels_indices[label])
        for i in range(parts):
            print(i)
            part_features = []
            part_labels = []
            for label in labels_indices:
                part_features += self.features[labels_indices[label][:(i + 1) * elements]].tolist()
                part_labels += [label] * elements
            PartialDataset(
                self,
                f"{i}",
                np.array(part_features),
                np.array(part_labels)
            ).save_dataset()

    def write_report(self):
        doc = Document()
        doc.add_heading(f"{language['dataset_heading']}: {self.dataset_name}", level=0)
        table = insert_table(doc, 'dataset_parameters')
        insert_param_to_table(table, 'bbox', self.bbox.format_osm())
        insert_param_to_table(table, 'sentinel2_name', [image.name for image in self.sentinel2image])
        insert_param_to_table(table, 'sentinel2_date', self.sentinel2image[0].date_time)
        insert_param_to_table(table, 'window_size', self.window_size * 2 + 1)
        if self.window_size > 0:
            insert_param_to_table(table, 'is_conv', self.is_conv)
            insert_param_to_table(table, 'rotations_and_reflections', self.apply_rotations_and_reflections)

        insert_param_to_table(table, 'features_count', self.features.shape[0])
        insert_param_to_table(table, 'features_params', self.features.shape[-1])

        doc.add_heading(language['labels_count'], level=1)
        table = doc.add_table(rows=1, cols=2)
        header_cells = table.rows[0].cells
        header_cells[0].text = language['class']
        header_cells[1].text = language['count']

        for i in range(self.labels.max() + 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(np.sum(self.labels == i))

        doc.save(os.path.join(workspace, "reports", "datasets", f"{self.dataset_name}.docx"))


class PointDataset:
    def __init__(
            self,
            dataset_name,
            label_source,
            class_column,
            sentinel2image,
            apply_cloud_mask=False,
            window_size=1,
            is_conv=False,
            apply_rotations_and_reflections=False
    ):
        if isinstance(dataset_name, str):
            self.dataset_name = dataset_name
        else:
            Exception("dataset_name must be str")
        self.label_source = label_source
        if isinstance(class_column, str):
            self.class_column = class_column
        else:
            Exception("class_column must be str")
        self.bbox = from_shapely_bounds(self.label_source.total_bounds, label_source.crs)
        if isinstance(sentinel2image, Sentinel2Image):
            self.sentinel2image = [sentinel2image]
        elif isinstance(sentinel2image, list):
            self.sentinel2image = sentinel2image
        else:
            Exception("identifier must be str")
        if isinstance(apply_cloud_mask, list):
            self.apply_cloud_mask = 0
            for el in apply_cloud_mask:
                self.apply_cloud_mask += el
        elif not apply_cloud_mask:
            self.apply_cloud_mask = apply_cloud_mask
        else:
            Exception("apply_cloud_mask must be False or list")
        if isinstance(window_size, int):
            self.window_size = window_size
        else:
            Exception("windows_size must be int")
        if isinstance(is_conv, bool):
            self.is_conv = is_conv
        else:
            Exception("ignore_off_scope must be bool")
        if self.window_size == 0:
            self.apply_rotations_and_reflections = False
        elif isinstance(apply_rotations_and_reflections, bool):
            self.apply_rotations_and_reflections = apply_rotations_and_reflections
        else:
            Exception("ignore_off_scope must be bool")

        snippets = []
        for i, row in label_source.iterrows():
            geom_bbox = Bbox(
                minx=(row.geometry.x - self.window_size * 10) // 60 * 60,
                miny=(row.geometry.y - self.window_size * 10) // 60 * 60,
                maxx=math.ceil((row.geometry.x + self.window_size * 10) / 60) * 60,
                maxy=math.ceil((row.geometry.y + self.window_size * 10) / 60) * 60,
                crs=label_source.crs
            )
            snippets.append([
                geom_bbox,
                [int((geom_bbox.maxy - row.geometry.y)//10), int((row.geometry.x - geom_bbox.minx)//10)],
                row[class_column],
                np.zeros((10, int((geom_bbox.maxy-geom_bbox.miny)//10), int((geom_bbox.maxx-geom_bbox.minx)//10)), dtype=np.uint16),
                [row.geometry.x, row.geometry.y]
            ])
        self.features = []
        self.labels = []
        if isinstance(self.apply_cloud_mask, int):
            for image in self.sentinel2image:
                with rasterio.open(
                        os.path.join(
                            workspace,
                            "cloud_masks",
                            f"{image.name.rsplit('_', 1)[0].replace('MSIL2A', 'MSIL1C')}.tif"
                        )) as mask:
                    for i in range(len(snippets)):
                        try:
                            masked = [el for el in mask.sample([snippets[i][-1]])]
                        except:
                            continue
                        if not (masked[0] & self.apply_cloud_mask) == 0:
                            snippets[i][1] = -1
        for i in range(len(snippets) - 1, -1, -1):
            if snippets[i][1] == -1:
                snippets = snippets[:i] + snippets[i + 1:]
        for image in self.sentinel2image:
            image.unzip()
            idx = 0
            layers = [
                el for el in os.listdir(os.path.join(workspace, "temp", image.identifier))
                if '60m' not in el and 'B01' not in el and el[-4:] == '.jp2'
            ]
            print(layers)
            for file_name in layers:
                with rasterio.open(os.path.join(workspace, "temp", image.identifier, file_name)) as channel:
                    for i in range(len(snippets)):
                        try:
                            masked, transform = rasterio.mask.mask(
                                channel,
                                [snippets[i][0].as_shapely()],
                                crop=True
                            )
                        except:
                            continue
                        xshift = int((transform[2] - snippets[i][0].minx) // 10)
                        yshift = int((snippets[i][0].maxy - transform[5]) // 10)
                        if '20m' in file_name:
                            if 'B01' in file_name:
                                continue
                            masked2 = np.zeros((1, masked.shape[1] * 2, masked.shape[2] * 2), dtype=masked.dtype)
                            for k in range(2):
                                for j in range(2):
                                    masked2[:, k::2, j::2] = masked
                            masked = masked2
                        elif '60m' in file_name:
                            continue
                        snippets[i][3][idx, yshift:masked.shape[1] + yshift, xshift:masked.shape[2] + xshift] = masked[
                            0]
                    idx += 1
        self.features = np.array([
            el[3][
                :,
                el[1][0] - self.window_size:el[1][0] + self.window_size + 1,
                el[1][1] - self.window_size:el[1][1] + self.window_size + 1
            ] for el in snippets
        ])
        self.labels = np.array([int(el[2]) for el in snippets])
        del label_source
        del snippets
        print(set(self.labels))
        print(self.labels.shape)
        print(self.features.shape)
        if self.apply_rotations_and_reflections:
            self.features = rotations_and_reflections(self.features)
            self.labels = np.block([
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels
            ])
        if not self.is_conv:
            self.features = self.features.reshape(
                self.features.shape[0],
                (self.window_size * 2 + 1) ** 2 * 10
            )
        if self.is_conv:
            self.features = self.features.transpose(0, 2, 3, 1)
        print(self.features.shape)
        self.write_report()

    def save_dataset(self):
        pickle.dump(
            self,
            open(os.path.join(workspace, "datasets", self.dataset_name), "wb"),
            protocol=pickle.HIGHEST_PROTOCOL
        )

    def split_dataset(self, parts, elements):
        indices = np.arange(len(self.labels))
        labels_indices = {el: indices[self.labels == el] for el in set(self.labels)}
        for label in labels_indices:
            np.random.shuffle(labels_indices[label])
        for i in range(parts):
            print(i)
            part_features = []
            part_labels = []
            for label in labels_indices:
                part_features += self.features[labels_indices[label][:(i + 1) * elements]].tolist()
                part_labels += [label] * elements
            PartialDataset(
                self,
                f"{i}",
                np.array(part_features),
                np.array(part_labels)
            ).save_dataset()

    def write_report(self):
        doc = Document()
        doc.add_heading(f"{language['dataset_heading']}: {self.dataset_name}", level=0)
        table = insert_table(doc, 'dataset_parameters')
        insert_param_to_table(table, 'bbox', self.bbox.format_osm())
        insert_param_to_table(table, 'sentinel2_name', [image.name for image in self.sentinel2image])
        insert_param_to_table(table, 'sentinel2_date', self.sentinel2image[0].date_time)
        insert_param_to_table(table, 'window_size', self.window_size * 2 + 1)
        if self.window_size > 0:
            insert_param_to_table(table, 'is_conv', self.is_conv)
            insert_param_to_table(table, 'rotations_and_reflections', self.apply_rotations_and_reflections)

        insert_param_to_table(table, 'features_count', self.features.shape[0])
        insert_param_to_table(table, 'features_params', self.features.shape[-1])

        doc.add_heading(language['labels_count'], level=1)
        table = doc.add_table(rows=1, cols=2)
        header_cells = table.rows[0].cells
        header_cells[0].text = language['class']
        header_cells[1].text = language['count']

        for i in range(self.labels.max() + 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(np.sum(self.labels == i))

        doc.save(os.path.join(workspace, "reports", "datasets", f"{self.dataset_name}.docx"))


class RasterDataset:
    def __init__(
            self,
            dataset_name,
            bbox,
            label_source,
            sentinel2image,
            apply_cloud_mask=False,
            window_size=1,
            is_conv=False,
            apply_rotations_and_reflections=False,
            no_data_value=6
    ):
        if isinstance(dataset_name, str):
            self.dataset_name = dataset_name
        else:
            Exception("dataset_name must be str")
        if isinstance(bbox, Bbox):
            self.bbox = bbox
        else:
            Exception("bbox must be Bbox")
        self.label_source = label_source
        if isinstance(sentinel2image, Sentinel2Image):
            self.sentinel2image = [sentinel2image]
        elif isinstance(sentinel2image, list):
            self.sentinel2image = sentinel2image
        else:
            Exception("identifier must be str")
        if isinstance(apply_cloud_mask, list):
            self.apply_cloud_mask = 0
            for el in apply_cloud_mask:
                self.apply_cloud_mask += el
        elif not apply_cloud_mask:
            self.apply_cloud_mask = apply_cloud_mask
        else:
            Exception("apply_cloud_mask must be False or list")
        if isinstance(window_size, int):
            self.window_size = window_size
        else:
            Exception("windows_size must be int")
        if isinstance(is_conv, bool):
            self.is_conv = is_conv
        else:
            Exception("ignore_off_scope must be bool")
        if self.window_size == 0:
            self.apply_rotations_and_reflections = False
        elif isinstance(apply_rotations_and_reflections, bool):
            self.apply_rotations_and_reflections = apply_rotations_and_reflections
        else:
            Exception("ignore_off_scope must be bool")

        self.height = label_source.shape[0]
        self.width = label_source.shape[1]
        self.features = []
        self.labels = []
        data = np.zeros((10, self.height, self.width), dtype=np.uint16)
        print(data.shape)
        for image in self.sentinel2image:
            image.unzip()

            idx = 0
            layers = [
                el for el in os.listdir(os.path.join(workspace, "temp", image.identifier))
                if '60m' not in el and 'B01' not in el and el[-4:] == '.jp2'
            ]
            print(layers)
            for file_name in layers:
                with rasterio.open(os.path.join(workspace, "temp", image.identifier, file_name)) as channel:
                    masked, transform = rasterio.mask.mask(channel, [bbox.as_shapely()], crop=True)
                    xshift = int((transform[2] - bbox.minx) // 10)
                    yshift = int((bbox.maxy - transform[5]) // 10)
                    if '20m' in file_name:
                        if 'B01' in file_name:
                            continue
                        masked2 = np.zeros((1, masked.shape[1] * 2, masked.shape[2] * 2), dtype=masked.dtype)
                        for i in range(2):
                            for j in range(2):
                                masked2[:, i::2, j::2] = masked
                        masked = masked2
                    elif '60m' in file_name:
                        continue
                    data[idx, yshift:masked.shape[1] + yshift, xshift:masked.shape[2] + xshift] = masked[0]
                    idx += 1
        cloud_mask = np.zeros((self.height, self.width), dtype=bool)
        if isinstance(self.apply_cloud_mask, int):
            for image in self.sentinel2image:
                with rasterio.open(
                        os.path.join(
                            workspace,
                            "cloud_masks",
                            f"{image.name.rsplit('_', 1)[0].replace('MSIL2A', 'MSIL1C')}.tif"
                        )) as mask:
                    masked, transform = rasterio.mask.mask(mask, [bbox.as_shapely()], crop=True)
                    xshift = int((transform[2] - bbox.minx) // 10)
                    yshift = int((bbox.maxy - transform[5]) // 10)
                    masked2 = np.zeros((1, masked.shape[1] * 6, masked.shape[2] * 6), dtype=masked.dtype)
                    for i in range(6):
                        for j in range(6):
                            masked2[:, i::6, j::6] = masked
                    masked = masked2
                    cloud_mask[
                        yshift:masked.shape[1] + yshift,
                        xshift:masked.shape[2] + xshift
                    ] |= (masked[0] & self.apply_cloud_mask) > 0
        if self.window_size > 0:
            data2 = np.zeros(
                (10, self.height + 2 * self.window_size, self.width + 2 * self.window_size),
                dtype=data.dtype
            )
            data2[
                :,
                self.window_size:data.shape[1] + self.window_size,
                self.window_size:data.shape[2] + self.window_size
            ] = data
            data = data2
        for i in range(self.height):
            for j in range(self.width):
                if not cloud_mask[i, j] and label_source[i, j] < no_data_value:
                    self.features.append(
                        data[
                        :,
                        i:i + 2 * self.window_size + 1,
                        j:j + 2 * self.window_size + 1
                        ])
                    self.labels.append(label_source[i, j])
        self.features = np.array(self.features)
        self.labels = np.array(self.labels)
        del label_source
        del data
        print(set(self.labels))
        print(self.labels.shape)
        print(self.features.shape)
        if self.apply_rotations_and_reflections:
            self.features = rotations_and_reflections(self.features)
            self.labels = np.block([
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels,
                self.labels
            ])
        if not self.is_conv:
            self.features = self.features.reshape(
                self.features.shape[0],
                (self.window_size * 2 + 1) ** 2 * 10
            )
        if self.is_conv:
            self.features = self.features.transpose(0, 2, 3, 1)
        print(self.features.shape)
        self.write_report()

    def save_dataset(self):
        pickle.dump(
            self,
            open(os.path.join(workspace, "datasets", self.dataset_name), "wb"),
            protocol=pickle.HIGHEST_PROTOCOL
        )

    def split_dataset(self, parts, elements):
        indices = np.arange(len(self.labels))
        labels_indices = {el: indices[self.labels == el] for el in set(self.labels)}
        for label in labels_indices:
            np.random.shuffle(labels_indices[label])
        for i in range(parts):
            print(i)
            part_features = []
            part_labels = []
            for label in labels_indices:
                part_features += self.features[labels_indices[label][:(i + 1) * elements]].tolist()
                part_labels += [label] * elements
            PartialDataset(
                self,
                f"{i}",
                np.array(part_features),
                np.array(part_labels)
            ).save_dataset()

    def write_report(self):
        doc = Document()
        doc.add_heading(f"{language['dataset_heading']}: {self.dataset_name}", level=0)
        table = insert_table(doc, 'dataset_parameters')
        insert_param_to_table(table, 'bbox', self.bbox.format_osm())
        insert_param_to_table(table, 'sentinel2_name', [image.name for image in self.sentinel2image])
        insert_param_to_table(table, 'sentinel2_date', self.sentinel2image[0].date_time)
        insert_param_to_table(table, 'window_size', self.window_size * 2 + 1)
        if self.window_size > 0:
            insert_param_to_table(table, 'is_conv', self.is_conv)
            insert_param_to_table(table, 'rotations_and_reflections', self.apply_rotations_and_reflections)

        insert_param_to_table(table, 'features_count', self.features.shape[0])
        insert_param_to_table(table, 'features_params', self.features.shape[-1])

        doc.add_heading(language['labels_count'], level=1)
        table = doc.add_table(rows=1, cols=2)
        header_cells = table.rows[0].cells
        header_cells[0].text = language['class']
        header_cells[1].text = language['count']

        for i in range(self.labels.max() + 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(np.sum(self.labels == i))

        doc.save(os.path.join(workspace, "reports", "datasets", f"{self.dataset_name}.docx"))


def process_edge_parcel(target, j, parcel, sheet, feature_path):
    image = rasterio.open(target)
    parcel_image, parcel_transform = rasterio.mask.mask(
        image,
        [parcel.geometry.buffer(0).intersection(sheet.geometry)],
        crop=True
    )
    key = str(j)
    parcel_meta = image.meta
    parcel_meta.update({"height": parcel_image.shape[1],
                        "width": parcel_image.shape[2],
                        "transform": parcel_transform})
    parcel_path = os.path.join(feature_path, "partial", key)
    if not os.path.exists(parcel_path):
        os.mkdir(parcel_path)
    parcel_path = os.path.join(parcel_path, f"{sheet['godlo']}_{parcel['class']}.tif")
    with rasterio.open(parcel_path, "w", **parcel_meta) as dest:
        dest.write(parcel_image)


def process_inside_parcel(target, j, parcel, sheet, feature_path, pyramid_base):
    image = rasterio.open(target)
    parcel_image, parcel_transform = rasterio.mask.mask(
        image,
        [parcel.geometry.buffer(0).intersection(sheet.geometry)],
        crop=True
    )
    key = str(j)
    level, parcel_image = pyramid_fit(parcel_image, pyramid_base)
    parcel_meta = image.meta
    parcel_meta.update({"height": parcel_image.shape[1],
                        "width": parcel_image.shape[2],
                        "transform": parcel_transform})
    level_path = os.path.join(feature_path, str(level))
    if not os.path.exists(level_path):
        os.mkdir(level_path)
    parcel_path = os.path.join(level_path, f"{key}_{parcel['class']}.tif")
    with rasterio.open(parcel_path, "w", **parcel_meta) as dest:
        dest.write(parcel_image)


def process_sheet(sheet, target, parcels, feature_path, pyramid_base):
    print(sheet['godlo'])
    if not os.path.exists(target):
        download_image(sheet['url_do_pob'], target)
    edge_parcels = parcels[parcels.intersects(sheet.geometry) & ~parcels.within(sheet.geometry)]
    inside_parcels = parcels[parcels.within(sheet.geometry)]
    print(edge_parcels['class'].sum(), len(edge_parcels))
    print(inside_parcels['class'].sum(), len(inside_parcels))
    print("Extracting dataset")
    with Pool() as pool:
        pool.starmap(
            process_edge_parcel,
            [(target, j, parcel, sheet, feature_path) for j, parcel in edge_parcels.iterrows()]
        )
        pool.starmap(
            process_inside_parcel,
            [(target, j, parcel, sheet, feature_path, pyramid_base) for j, parcel in inside_parcels.iterrows()]
        )
    print("Dataset extracted")


def merge_partial_parcels(feature_path, el, pyramid_base):
    dir_path = os.path.join(feature_path, "partial", el)
    with rasterio.open(os.path.join(dir_path, os.listdir(dir_path)[0]), 'r') as image:
        merge_meta = image.meta
    merge_image, merge_transform = rasterio.merge.merge([os.path.join(dir_path, el2) for el2 in os.listdir(dir_path)])
    label = int(1 in [el2.rsplit("_", 1)[-1] for el2 in os.listdir(dir_path)])
    level, merge_image = pyramid_fit(merge_image, pyramid_base)
    merge_meta.update({"height": merge_image.shape[1],
                       "width": merge_image.shape[2],
                       "transform": merge_transform})
    level_path = os.path.join(feature_path, str(level))
    if not os.path.exists(level_path):
        os.mkdir(level_path)
    merge_path = os.path.join(level_path, f"{el}_{label}.tif")
    with rasterio.open(merge_path, "w", **merge_meta) as dest:
        dest.write(merge_image)
    shutil.rmtree(dir_path)


def pyramid_fit(data, pyramid_base):
    height = data.shape[1]
    width = data.shape[2]
    level = max(0, int(math.ceil(math.log(max(height, width) / pyramid_base, 2))))
    new_data = np.zeros((data.shape[0], pyramid_base * 2 ** level, pyramid_base * 2 ** level), dtype=data.dtype)
    new_data[:, :height, :width] = data
    return level, new_data


class CadastralDataset:
    def __init__(self, name, pyramid_base=150):
        sheets = gpd.read_file("C:/Users/trole/Documents/Łódź/arkusze2021Łódź.shp")
        sheets = sheets[sheets['piksel'] == 0.05]
        parcels = gpd.read_file("C:/Users/trole/Documents/Łódź/etykietyBezNisko.shp")
        self.name = name
        self.feature_path = os.path.join(workspace, "parcels", name)
        self.pyramid_base = pyramid_base
        if not os.path.exists(self.feature_path):
            os.mkdir(self.feature_path)
            self.feature_path = os.path.join(self.feature_path, "unnormalized")
            os.mkdir(self.feature_path)
            os.mkdir(os.path.join(self.feature_path, "partial"))
        else:
            self.feature_path = os.path.join(self.feature_path, "unnormalized")
        for i, sheet in sheets.iterrows():
            if parcels.intersects(sheet.geometry).any():
                process_sheet(
                    sheet,
                    f"temp/ortofoto/{sheet['godlo']}.tif",
                    parcels,
                    self.feature_path,
                    self.pyramid_base
                )
        with Pool() as pool:
            pool.starmap(merge_partial_parcels, [
                (
                    self.feature_path,
                    el,
                    self.pyramid_base
                ) for el in os.listdir(os.path.join(self.feature_path, "partial"))
            ])
        print(len(os.listdir(self.feature_path)))


def create_test_dataset(dataset_name, class_raster, bbox, mode, count, included_values):
    with rasterio.open(class_raster) as image:
        data = image.read(1)
        ll_corner = image.transform * (0, 0)
        if bbox is None:
            xmin = 0
            xmax = image.height
            ymin = 0
            ymax = image.width
        else:
            bbox = bbox.to_crs(image.crs)
            xmin = int((ll_corner[1] - bbox.maxy) // 10)
            ymin = int((bbox.minx - ll_corner[0]) // 10)
            xmax = int((ll_corner[1] - bbox.miny) // 10)
            ymax = int((bbox.maxx - ll_corner[0]) // 10)
        data = data[xmin:xmax, ymin:ymax]
        if mode == "total_random":
            arr = random.sample(np.array(np.where(data == data)).transpose().tolist(), count)
        elif mode == "proportional":
            unique, unique_counts = np.unique(data, return_counts=True)
            unique = list(unique)
            unique_counts = list(unique_counts)
            for i in range(len(unique) - 1, -1, -1):
                if unique[i] not in included_values:
                    unique = unique[:i] + unique[i + 1:]
                    unique_counts = unique_counts[:i] + unique_counts[i + 1:]
            unique = np.array(unique)
            unique_counts = np.array(unique_counts)
            prop = count * unique_counts / data.shape[0]
            base = prop.astype(np.int32)
            rest = prop - base
            while np.sum(base) < count:
                max_idx = np.argmax(rest)
                base[max_idx] += 1
                rest[max_idx] = 0
            arr = []
            for i in range(unique.shape[0]):
                arr += random.sample(np.array(np.where(data == unique[i])).transpose().tolist(), base[i])
        elif mode == "equal":
            unique = np.array(included_values)
            base = count // unique.shape[0] * np.ones(unique.shape, dtype=np.int16)
            rest = count % unique.shape[0]
            base[random.sample(range(unique.shape[0]), rest)] += 1
            arr = []
            for i in range(unique.shape[0]):
                arr += random.sample(np.array(np.where(data == unique[i])).transpose().tolist(), base[i])
        arr = np.array(arr)
        arr = np.array([ll_corner[1] - (arr[:, 0] + xmin) * 10 - 5, (arr[:, 1] + ymin) * 10 + ll_corner[0] + 5,
                        data[arr[:, 0], arr[:, 1]], -np.ones(arr.shape[0])])
        df = pd.DataFrame(data=arr.transpose(), columns=["y", "x", "value", "ground_truth"])
        gdf = gpd.GeoDataFrame(df, geometry=gpd.points_from_xy(df.x, df.y), crs=image.crs)
        gdf.to_file(os.path.join(workspace, "test_datasets", f"{dataset_name}.geojson"))


def load_dataset(dataset_name):
    return pickle.load(open(os.path.join(workspace, "datasets", dataset_name), "rb"))


def rotations_and_reflections(dataset):
    dataset2 = np.rot90(dataset, axes=(1, 2))
    dataset3 = np.rot90(dataset2, axes=(1, 2))
    dataset4 = np.rot90(dataset3, axes=(1, 2))
    dataset5 = dataset[:, :, ::-1]
    dataset6 = dataset2[:, :, ::-1]
    dataset7 = dataset3[:, :, ::-1]
    dataset8 = dataset4[:, :, ::-1]
    return np.block(
        [[[[dataset]]], [[[dataset2]]], [[[dataset3]]], [[[dataset4]]], [[[dataset5]]], [[[dataset6]]], [[[dataset7]]],
         [[[dataset8]]]])
